#!/usr/bin/env python
# coding: utf-8

import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import threading
import requests
from bs4 import BeautifulSoup
import urllib.parse
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from PIL import Image
from docx.oxml.ns import qn
import os
from VGH_function import *
from datetime import datetime, timedelta
import time
import random
def set_font_size(doc, size):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(size)

def set_paragraph_spacing(doc, spacing=0):
    """Set paragraph spacing for all paragraphs in the document."""
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = Pt(spacing)
        paragraph.paragraph_format.space_before = Pt(spacing)
        paragraph.paragraph_format.space_after = Pt(spacing)

def add_table(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # 設置表頭的字體大小
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = str(column_name)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(6)
            paragraph.paragraph_format.line_spacing = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    last_paragraph = doc.paragraphs[-1]
    last_paragraph.paragraph_format.space_after = Pt(0)
    last_paragraph.paragraph_format.space_before = Pt(0)
    last_paragraph.paragraph_format.line_spacing = Pt(0)
    
    # 添加數據行
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6)
                paragraph.paragraph_format.line_spacing = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                
    for col in table.columns:
        max_length = max(len(cell.text) for cell in col.cells)
        col_width = Inches(max_length)
        for cell in col.cells:
            cell.width = col_width

def convert_date(date_str):
    date_str = date_str[3:8]
    return date_str

def convert_drug(data_drug):
    data_drug = data_drug.split(" ")[:2]
    data_drug = " ".join(data_drug)
    return data_drug

def convert_drug_date(data_drug_date):
    data_drug_date = data_drug_date[5:10]
    return data_drug_date

class VGHLogin:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        self.csrf_token = None
        self.base_url = "https://eip.vghtpe.gov.tw/login.php"
    
    def get_login_page(self):
        """取得登入頁面並解析CSRF token"""
        try:
            response = self.session.get(self.base_url)
            response.raise_for_status()
            
            # 解析HTML取得CSRF token
            soup = BeautifulSoup(response.text, 'html.parser')
            csrf_meta = soup.find('meta', {'name': 'csrf-token'})
            if csrf_meta:
                self.csrf_token = csrf_meta.get('content')
            
            return True
        except requests.RequestException as e:
            print(f"取得登入頁面失敗: {e}")
            return False
    
    def login(self, username, password):
        """執行登入"""
        if not self.get_login_page():
            return False
        
        # 準備登入資料
        login_data = {
            'login_name': username,
            'password': password,
            'loginCheck': '1',
            'fromAjax': '1'
        }
        
        # 設定headers
        headers = {
            'X-CSRF-TOKEN': self.csrf_token,
            'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
            'X-Requested-With': 'XMLHttpRequest',
            'Referer': self.base_url
        }
        
        try:
            # 發送登入請求
            login_url = urllib.parse.urljoin(self.base_url, '/login_action.php')
            response = self.session.post(
                login_url,
                data=login_data,
                headers=headers
            )
            response.raise_for_status()
            
            # 解析回應
            result = response.json()

            if 'error' in result:
                error_code = int(result['error'])
                if error_code == 0:
                    if 'url' in result:
                        dashboard_response = self.session.get("https://eip.vghtpe.gov.tw/"+result['url'])
                        login_url="https://eip.vghtpe.gov.tw/"+dashboard_response.text.split("/")[1][:-2]
                        dashboard_response = self.session.get(login_url)
                        return True
                else:
                    return False
            else:
                return False
                
        except requests.RequestException as e:
            return False
        except ValueError as e:
            return False
    
    def get_page_after_login(self, url):
        """登入後取得其他頁面"""
        try:
            response = self.session.get(url)
            response.raise_for_status()
            return response.text
        except requests.RequestException as e:
            return None

    def get_img_after_login(self, url):
        """登入後取得其他頁面"""
        try:
            response = self.session.get(url)
            response.raise_for_status()
            return response
        except requests.RequestException as e:
            return None



class VGHPatientGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("查房摘要產生器")
        self.root.geometry("1080x900")
        self.root.configure(bg='#f0f0f0')
        
        # 設定樣式
        style = ttk.Style()
        style.theme_use('clam')
        
        self.vgh = None
        self.pat_data = []
        self.search_type = ""
        self.docID = ""
        
        self.create_widgets()
    
    def create_widgets(self):
        # 主標題
        title_frame = tk.Frame(self.root, bg='#2c3e50', height=80)
        title_frame.pack(fill='x', padx=10, pady=(10, 0))
        title_frame.pack_propagate(False)
        
        title_label = tk.Label(title_frame, text="VGH 查房摘要產生器", 
                              font=('Arial', 18, 'bold'), fg='white', bg='#2c3e50')
        title_label.pack(expand=True)
        
        # 說明文字
        info_frame = tk.Frame(self.root, bg='#ecf0f1', relief='solid', bd=1)
        info_frame.pack(fill='x', padx=10, pady=10)
        
        info_text = """
                        此程式可以自動查詢病人資料，製作一份WORD的查房摘要
                        ***注意：若跳出密碼到期，請先更改密碼後再使用此程式
                        ***注意：若病人太多可能會被資訊室鎖住該台電腦一陣子，可以重新開機後稍等一下
                        作者的燈號為: 8375K，如果有任何問題或建議，歡迎聯絡!!!"""
        
        info_label = tk.Label(info_frame, text=info_text, font=('Arial', 10), 
                             bg='#ecf0f1', fg='#2c3e50', justify='left')
        info_label.pack(padx=15, pady=10)
        
        # 登入區域
        login_frame = tk.LabelFrame(self.root, text="登入資訊", font=('Arial', 12, 'bold'),
                                   fg='#2c3e50', bg='#f0f0f0')
        login_frame.pack(fill='x', padx=10, pady=5)
        
        # 帳號輸入
        tk.Label(login_frame, text="帳號:", font=('Arial', 11), bg='#f0f0f0').grid(row=0, column=0, padx=10, pady=5, sticky='w')
        self.username_entry = tk.Entry(login_frame, font=('Arial', 11), width=20)
        self.username_entry.grid(row=0, column=1, padx=10, pady=5, sticky='w')
        
        # 密碼輸入
        tk.Label(login_frame, text="密碼:", font=('Arial', 11), bg='#f0f0f0').grid(row=1, column=0, padx=10, pady=5, sticky='w')
        self.password_entry = tk.Entry(login_frame, show="*", font=('Arial', 11), width=20)
        self.password_entry.grid(row=1, column=1, padx=10, pady=5, sticky='w')
        
        # 登入按鈕
        self.login_btn = tk.Button(login_frame, text="登入", font=('Arial', 11, 'bold'),
                                  bg='#3498db', fg='white', width=10, height=1,
                                  command=self.login_action)
        self.login_btn.grid(row=0, column=2, rowspan=2, padx=20, pady=5)
        
        # 登入狀態顯示
        self.login_status = tk.Label(login_frame, text="未登入", font=('Arial', 10),
                                    fg='red', bg='#f0f0f0')
        self.login_status.grid(row=2, column=0, columnspan=3, pady=5)
        
        # 搜尋選項區域
        search_frame = tk.LabelFrame(self.root, text="搜尋選項", font=('Arial', 12, 'bold'),
                                    fg='#2c3e50', bg='#f0f0f0')
        search_frame.pack(fill='x', padx=10, pady=5)
        
        # 搜尋類型選擇
        self.search_var = tk.StringVar(value="doc")
        tk.Radiobutton(search_frame, text="依燈號搜尋", variable=self.search_var, value="doc",
                      font=('Arial', 11), bg='#f0f0f0', command=self.update_search_input).grid(row=0, column=0, padx=10, pady=5, sticky='w')
        tk.Radiobutton(search_frame, text="依病房搜尋", variable=self.search_var, value="ward",
                      font=('Arial', 11), bg='#f0f0f0', command=self.update_search_input).grid(row=0, column=1, padx=10, pady=5, sticky='w')
        tk.Radiobutton(search_frame, text="依病歷號搜尋", variable=self.search_var, value="pat",
                      font=('Arial', 11), bg='#f0f0f0', command=self.update_search_input).grid(row=0, column=2, padx=10, pady=5, sticky='w')
        
        # 搜尋輸入區域
        input_frame = tk.Frame(search_frame, bg='#f0f0f0')
        input_frame.grid(row=1, column=0, columnspan=4, sticky='ew', padx=10, pady=5)
        
        self.input_label = tk.Label(input_frame, text="請輸入燈號(四碼):", font=('Arial', 11), bg='#f0f0f0')
        self.input_label.pack(side='left')
        
        self.search_entry = tk.Entry(input_frame, font=('Arial', 11), width=30)
        self.search_entry.pack(side='left', padx=10)
        
        # 病歷號多筆輸入區域（初始隱藏）
        self.multi_input_frame = tk.Frame(search_frame, bg='#f0f0f0')
        
        tk.Label(self.multi_input_frame, text="病歷號列表（每行一個）:", font=('Arial', 11), bg='#f0f0f0').pack(anchor='w')
        self.multi_entry = scrolledtext.ScrolledText(self.multi_input_frame, width=50, height=6, font=('Arial', 10))
        self.multi_entry.pack(pady=5)
        
        # 搜尋按鈕
        self.search_btn = tk.Button(search_frame, text="開始搜尋", font=('Arial', 11, 'bold'),
                                   bg='#27ae60', fg='white', width=12, height=1,
                                   command=self.search_patients)
        self.search_btn.grid(row=0, column=5, columnspan=4, pady=10)

                # ===================== 護理紀錄 VIPFSN 區塊 =====================
        vip_frame = tk.LabelFrame(self.root,
                                  text="護理紀錄 V / I / P / ECMO / F / S / N",
                                  font=('Arial', 12, 'bold'),
                                  fg='#2c3e50',
                                  bg='#f0f0f0')
        vip_frame.pack(fill='x', padx=10, pady=(0, 5))

        tk.Label(vip_frame, text="病歷號:", font=('Arial', 11), bg='#f0f0f0')\
            .grid(row=0, column=0, padx=10, pady=5, sticky='w')

        # 單一病歷號輸入
        self.vipfsn_entry = tk.Entry(vip_frame, font=('Arial', 11), width=20)
        self.vipfsn_entry.grid(row=0, column=1, padx=10, pady=5, sticky='w')

        # 顯示 VIPFSN 按鈕
        self.vipfsn_btn = tk.Button(
            vip_frame,
            text="顯示 VIPFSN",
            font=('Arial', 11, 'bold'),
            bg='#8e44ad',
            fg='white',
            width=12,
            height=1,
            command=self.show_vipfsn
        )
        self.vipfsn_btn.grid(row=0, column=2, padx=10, pady=5)

        # 狀態顯示（讀取中 / 成功 / 失敗）
        self.vipfsn_status = tk.Label(
            vip_frame,
            text="",
            font=('Arial', 10),
            bg='#f0f0f0',
            fg='#2c3e50'
        )
        self.vipfsn_status.grid(row=1, column=0, columnspan=3,
                                sticky='w', padx=10, pady=(0, 5))

        
        # 進階設定區域（可收合）
        
        self.advanced_expanded = tk.BooleanVar(value=False)
        advanced_header_frame = tk.Frame(self.root, bg='#f0f0f0')
        advanced_header_frame.pack(fill='x', padx=10, pady=(5, 0))
        
        self.advanced_toggle_btn = tk.Button(advanced_header_frame, text="▶ 進階設定", 
                                            font=('Arial', 11, 'bold'), fg='#2c3e50', bg='#f0f0f0',
                                            relief='flat', command=self.toggle_advanced_settings)
        self.advanced_toggle_btn.pack(side='left')
        
        # 進階設定內容區域
        self.advanced_frame = tk.LabelFrame(self.root, text="", font=('Arial', 12, 'bold'),
                                           fg='#2c3e50', bg='#f0f0f0', bd=1, relief='solid')
        self.advanced_frame.pack(fill='x', padx=10, pady=5)
        self.advanced_frame.pack_forget()

        # 建立進階設定內容
        self.create_advanced_settings()
        
        # 進度區域
        progress_frame = tk.LabelFrame(self.root, text="處理進度", font=('Arial', 12, 'bold'),
                                      fg='#2c3e50', bg='#f0f0f0')
        progress_frame.pack(fill='x', padx=10, pady=5)
        
        self.progress_var = tk.StringVar(value="等待開始...")
        self.progress_label = tk.Label(progress_frame, textvariable=self.progress_var,
                                      font=('Arial', 10), bg='#f0f0f0')
        self.progress_label.pack(pady=10)
        
        self.progress_bar = ttk.Progressbar(progress_frame, mode='indeterminate')
        self.progress_bar.pack(fill='x', padx=20, pady=(0, 10))
        
        # 按鈕區域
        button_frame = tk.Frame(self.root, bg='#f0f0f0')
        button_frame.pack(fill='x', padx=10, pady=5)
        
        self.generate_btn = tk.Button(button_frame, text="產生Word報告", font=('Arial', 11, 'bold'),
                                     bg='#e74c3c', fg='#f0f0f0', width=15, height=1,
                                     command=self.generate_report, state='disabled')
        self.generate_btn.pack(side='left', padx=5)

        tk.Button(button_frame, text="清除結果", font=('Arial', 11),
            bg='#95a5a6', fg='white', width=10, height=1,
            command=self.clear_results).pack(side='left', padx=5)
        
        # 結果區域（可收合）
        self.result_expanded = tk.BooleanVar(value=False)
        result_header_frame = tk.Frame(self.root, bg='#f0f0f0')
        result_header_frame.pack(fill='x', padx=10, pady=(5, 0))
        
        self.result_toggle_btn = tk.Button(result_header_frame, text="▶ 處理結果", 
                                          font=('Arial', 11, 'bold'), fg='#2c3e50', bg='#f0f0f0',
                                          relief='flat', command=self.toggle_result_area)
        self.result_toggle_btn.pack(side='left')
        
        # 結果內容區域
        self.result_frame = tk.LabelFrame(self.root, text="", font=('Arial', 12, 'bold'),
                                         fg='#2c3e50', bg='#f0f0f0', bd=1, relief='solid')
        self.result_frame.pack(fill='both', expand=True, padx=10, pady=(0, 5))
        
        self.result_text = scrolledtext.ScrolledText(self.result_frame, font=('Arial', 10), 
                                                    height=8, bg='white')
        self.result_text.pack(fill='x', expand=False, padx=10, pady=10)

    def create_advanced_settings(self):
        """建立進階設定內容"""
        # 全選/全不選按鈕
        control_frame = tk.Frame(self.advanced_frame, bg='#f0f0f0')
        control_frame.pack(fill='x', padx=10, pady=5)
        
        tk.Button(control_frame, text="全選", font=('Arial', 9), bg='#3498db', fg='white',
                 command=self.select_all_items, width=8).pack(side='left', padx=2)
        tk.Button(control_frame, text="全不選", font=('Arial', 9), bg='#e74c3c', fg='white',
                 command=self.deselect_all_items, width=8).pack(side='left', padx=2)
        tk.Button(control_frame, text="恢復預設", font=('Arial', 9), bg='#f39c12', fg='white',
                 command=self.reset_to_default, width=8).pack(side='left', padx=2)
        
        # 創建滾動區域
        canvas = tk.Canvas(self.advanced_frame, bg='#f0f0f0', height=200)
        scrollbar = ttk.Scrollbar(self.advanced_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='#f0f0f0')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 建立檢查項目變數字典
        self.lab_vars = {}
        
        # 生化項目
        biochem_frame = tk.LabelFrame(scrollable_frame, text="生化檢驗", font=('Arial', 11, 'bold'),
                                     bg='#f0f0f0', fg='#2c3e50')
        biochem_frame.pack(fill='x', padx=5, pady=5)
        
        biochem_items = [
            'TP', 'ALB', 'CA', 'CACAL', 'CHOL', 'BUN', 'UA', 'CREA', 'eGFR(EPI)', 'BILIT',
            'ALKP', 'LDH', 'ALT', 'AST', 'FIB-4', 'NA', 'K', 'CL', 'GLU', 'IP', 'CK', 'GGT',
            'HDLC', 'LDLC', 'LDL(計算值)', 'RISKF', 'TG', 'CO2', 'DBILI', 'IRON', 'TIBC',
            'CKMB', 'CKMB(POCT)', 'TROP', 'TROP(POCT)', 'CRP', 'hs-cTnt', 'CRP(POCT)',
            'PCT(POCT)', 'eGFR(M)', 'eGFR(PED)', 'eGFR(C-G)未BSA校正', 'eGFR(C-G)BSA校正',
            'Amylase', 'Lipase', 'Freekappa', 'Freelambda', 'kappa/lambda', 'ESR', 'lactate',
            'procalcitonin(PCT)', 'Free Ca++', 'Mg', 'Amonia(NH3)', 'NT-ProBNP', 'Intact-PTH',
            'Myoglobin', 'TSAT', 'Ferritin'
        ]
        
        self.create_checkbox_grid(biochem_frame, biochem_items, 'biochem')
        
        # 血球項目
        hematology_frame = tk.LabelFrame(scrollable_frame, text="血球檢驗", font=('Arial', 11, 'bold'),
                                        bg='#f0f0f0', fg='#2c3e50')
        hematology_frame.pack(fill='x', padx=5, pady=5)
        
        hematology_items = [
            'WBC', 'RBC', 'HGB', 'HCT', 'MCV', 'MCH', 'MCHC', 'RDW', 'PLT', 'INR(PT)',
            'PT', 'APTT', 'BAND', 'SEG', 'LYM', 'MONO', 'EOS', 'BASO', 'D-dimer',
            'FDP', 'FIBRINOGEN'
        ]
        
        self.create_checkbox_grid(hematology_frame, hematology_items, 'hematology')
        
        canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar.pack(side="right", fill="y", pady=10)
        self.reset_to_default()
        
    def create_checkbox_grid(self, parent, items, category):
        """建立複選框網格"""
        for i, item in enumerate(items):
            var = tk.BooleanVar(value=False)  # 預設全選
            self.lab_vars[f"{category}_{item}"] = var
            
            cb = tk.Checkbutton(parent, text=item, variable=var, font=('Arial', 9),
                               bg='#f0f0f0', anchor='w')
            row = i // 6  # 每行6個
            col = i % 6
            cb.grid(row=row, column=col, sticky='w', padx=5, pady=2)
    
    def toggle_advanced_settings(self):
        """切換進階設定區域顯示/隱藏"""
        if self.advanced_expanded.get():
            self.advanced_frame.pack_forget()
            self.advanced_toggle_btn.config(text="▶ 進階設定")
            self.advanced_expanded.set(False)
        else:
            self.advanced_frame.pack(fill='x', padx=10, pady=(0, 5))
            self.advanced_toggle_btn.config(text="▼ 進階設定")
            self.advanced_expanded.set(True)
        self.root.update_idletasks()
    
    def toggle_result_area(self):
        """切換結果區域顯示/隱藏"""
        if self.result_expanded.get():
            self.result_frame.pack_forget()
            self.result_toggle_btn.config(text="▶ 處理結果")
            self.result_expanded.set(False)
        else:
            self.result_frame.pack(fill='both', expand=True, padx=10, pady=(0, 5))
            self.result_toggle_btn.config(text="▼ 處理結果")
            self.result_expanded.set(True)
    
    def select_all_items(self):
        """全選所有檢驗項目"""
        for var in self.lab_vars.values():
            var.set(True)
    
    def deselect_all_items(self):
        """取消選擇所有檢驗項目"""
        for var in self.lab_vars.values():
            var.set(False)
    
    def reset_to_default(self):
        self.deselect_all_items()
        """恢復預設選擇（全選）"""
        Default=["NA", "K", "BUN", "CREA", "ALT", "BILIT", "CRP", "WBC","HGB","PLT","SEG","PT","APTT"]
        for key, var in self.lab_vars.items():
            key=key.split("_")[1]
            if key in Default:
                var.set(True)
    
    def get_selected_lab_items(self):
        """取得已選擇的檢驗項目"""
        selected = {}
        for key, var in self.lab_vars.items():
            if var.get():
                category, item = key.split('_', 1)
                if category not in selected:
                    selected[category] = []
                selected[category].append(item)
        return selected
        
        
    
    def update_search_input(self):
        """更新搜尋輸入界面"""
        search_type = self.search_var.get()
        
        if search_type == "doc":
            self.input_label.config(text="請輸入燈號(四碼):")
            self.multi_input_frame.grid_remove()
            self.search_entry.pack(side='left', padx=10)
        elif search_type == "ward":
            self.input_label.config(text="請輸入病房(Ex: A101):")
            self.multi_input_frame.grid_remove()
            self.search_entry.pack(side='left', padx=10)
        else:  # pat
            self.input_label.config(text="")
            self.search_entry.pack_forget()
            self.multi_input_frame.grid(row=1, column=0, columnspan=4, sticky='ew', padx=10, pady=5)
    
    def login_action(self):
        """登入處理"""
        username = self.username_entry.get().strip()
        password = self.password_entry.get().strip()
        
        if not username or not password:
            messagebox.showerror("錯誤", "請輸入帳號和密碼")
            return
        
        # 禁用登入按鈕，避免重複點擊
        self.login_btn.config(state='disabled', text="登入中...")
        self.login_status.config(text="登入中...", fg="orange")
        
        def login_thread():
            try:
                self.vgh = VGHLogin()
                if self.vgh.login(username, password):
                    self.root.after(0, lambda: self.login_success())
                else:
                    self.root.after(0, lambda: self.login_failed("登入失敗，請檢查帳號密碼"))
            except Exception as e:
                self.root.after(0, lambda: self.login_failed(f"登入錯誤: {str(e)}"))
        
        threading.Thread(target=login_thread, daemon=True).start()
    
    def login_success(self):
        """登入成功處理"""
        self.login_btn.config(state='normal', text="登入")
        self.login_status.config(text="✅ 登入成功", fg="green")
        self.result_text.insert(tk.END, "登入成功！可以開始搜尋病人資料。\n")
        self.result_text.see(tk.END)
    
    def login_failed(self, error_msg):
        """登入失敗處理"""
        self.login_btn.config(state='normal', text="登入")
        self.login_status.config(text="❌ 登入失敗", fg="red")
        messagebox.showerror("登入失敗", error_msg)
    
    def search_patients(self):
        """搜尋病人"""
        if not self.vgh:
            messagebox.showerror("錯誤", "請先登入系統")
            return
        
        search_type = self.search_var.get()
        
        if search_type in ["doc", "ward"]:
            search_value = self.search_entry.get().strip()
            if not search_value:
                messagebox.showerror("錯誤", "請輸入搜尋條件")
                return
        else:  # pat
            search_values = self.multi_entry.get("1.0", tk.END).strip().split('\n')
            search_values = [v.strip() for v in search_values if v.strip()]
            if not search_values:
                messagebox.showerror("錯誤", "請輸入至少一個病歷號")
                return
        
        # 禁用搜尋按鈕
        self.search_btn.config(state='disabled', text="搜尋中...")
        self.progress_bar.start()
        self.progress_var.set("正在搜尋病人資料...")
        
        def search_thread():
            try:
                if search_type == "doc":
                    self.docID = search_value
                    self.search_type = search_type
                    pat_data = get_searched_patient(self.vgh, ward="0", patID="", docID=search_value)
                elif search_type == "ward":
                    self.docID = ""
                    self.search_type = search_type
                    pat_data = get_searched_patient(self.vgh, ward=search_value, patID="", docID="")
                else:  # pat
                    self.docID = ""
                    self.search_type = search_type
                    pat_data = []
                    for patID in search_values:
                        try:
                            patient = get_searched_patient(self.vgh, ward="0", patID=patID, docID="")
                            if patient:
                                pat_data.append(patient[0])
                        except:
                            continue
                
                self.pat_data = pat_data
                self.root.after(0, lambda: self.search_complete(len(pat_data)))
                
            except Exception as e:
                self.root.after(0, lambda: self.search_failed(str(e)))
        
        threading.Thread(target=search_thread, daemon=True).start()

    def show_vipfsn_complete(self, histno, data):
        """成功取得 VIPFSN 後顯示在結果區"""
        self.vipfsn_btn.config(state='normal', text="顯示 VIPFSN")
        self.vipfsn_status.config(text=f"取得 {histno} 最新護理紀錄完成", fg="green")

        def val(d, key):
            return d.get(key, "") if isinstance(d, dict) else ""

        V = data.get("V", {})
        I = data.get("I", {})
        P = data.get("P", {})
        ECMO = data.get("ECMO", {})
        F = data.get("F", {})
        S = data.get("S", {})
        N = data.get("N", {})

        lines = []
        # V
        lines.append(
            "V (Ventilation): "
            f"mode={val(V,'mode')}, "
            f"FiO2={val(V,'FiO2')}, "
            f"PC={val(V,'PC')}, "
            f"PEEP={val(V,'PEEP')}"
        )

        # I
        lines.append(
            "I (Infection): "
            f"Abx={val(I,'abx')}, "
            f"Culture={val(I,'culture')}, "
            f"Infection status={val(I,'infection_status')}"
        )

        # P
        lines.append(
            "P (Pressure): "
            f"Perdipine={val(P,'Perdipine_ml_hr')} ml/hr, "
            f"NTG={val(P,'NTG_ml_hr')} ml/hr, "
            f"BP goal={val(P,'BP_goal')}, "
            f"Levophed={val(P,'Levophed_mcg_kg_min')} mcg/kg/min, "
            f"Epi={val(P,'Epinephrine_mcg_kg_min')} mcg/kg/min, "
            f"Dopamine={val(P,'Dopamine_mcg_kg_min')} mcg/kg/min, "
            f"Vasopressin={val(P,'Vasopressin_U_hr')} U/hr"
        )

        # ECMO
        lines.append(
            "ECMO / VAD / IABP: "
            f"Mode={val(ECMO,'mode')}, "
            f"FiO2={val(ECMO,'FiO2')}, "
            f"Gas flow={val(ECMO,'gas_flow')}, "
            f"Rate={val(ECMO,'rate')}, "
            f"VAD ratio={val(ECMO,'VAD_ratio')}, "
            f"IABP ratio={val(ECMO,'IABP_ratio')}"
        )

        # F
        lines.append(
            "F (Fluid): "
            f"CVVH={val(F,'CVVH')}, "
            f"HD={val(F,'HD')}, "
            f"Dry weight={val(F,'dry_weight')}, "
            f"Diuretics={val(F,'diuretics')}, "
            f"I/O={val(F,'IO_summary')}"
        )

        # S
        lines.append(
            "S (Sedation): "
            f"Dormicum={val(S,'DORMICUM')}, "
            f"Propofol={val(S,'PROPOFOL')}, "
            f"Fentanyl={val(S,'FENTANYL')}"
        )

        # N
        lines.append(f"N (Nutrition): Route={val(N,'route')}")

        text_block = "\n".join(lines)

        self.result_text.insert(
            tk.END,
            f"\n=== 病歷號 {histno} 最新護理紀錄 VIPFSN ===\n"
        )
        self.result_text.insert(tk.END, text_block + "\n")

        # ===== DEBUG：顯示護理紀錄原始文字的一小段 =====
        raw_text = data.get("raw_text", "")
        raw_len = len(raw_text)
        self.result_text.insert(
            tk.END,
            f"\n[DEBUG] 抓到的護理紀錄文字長度：{raw_len}\n")

        if raw_len > 0:
            preview = raw_text.replace("\r", " ").replace("\n", " ")
            preview = preview[:300]  # 只秀前 300 個字，避免太長
            self.result_text.insert(
                tk.END,
                f"[DEBUG] 護理紀錄前 300 字預覽：{preview}...\n")

        self.result_text.see(tk.END)




    def show_vipfsn_failed(self, error_msg):
        """取得 VIPFSN 失敗時處理"""
        self.vipfsn_btn.config(state='normal', text="顯示 VIPFSN")
        self.vipfsn_status.config(text="取得護理紀錄失敗", fg="red")
        self.result_text.insert(tk.END, f"\n❌ 取得護理紀錄失敗：{error_msg}\n")
        self.result_text.see(tk.END)
        messagebox.showerror("錯誤", f"取得護理紀錄失敗：{error_msg}")

    def show_vipfsn(self):
        """按下按鈕，依病歷號抓最新護理紀錄 VIPFSN"""
        if not self.vgh:
            messagebox.showerror("錯誤", "請先登入系統")
            return

        histno = self.vipfsn_entry.get().strip()
        if not histno:
            messagebox.showerror("錯誤", "請先輸入病歷號")
            return

        # 按鈕 & 狀態顯示更新
        self.vipfsn_btn.config(state='disabled', text="讀取中...")
        self.vipfsn_status.config(text=f"正在讀取 {histno} 最新護理紀錄...", fg="orange")
        self.result_text.insert(tk.END, f"\n▶ 開始讀取病歷號 {histno} 的最新護理紀錄...\n")
        self.result_text.see(tk.END)

        def worker():
            try:
                data = get_latest_nursing_data(self.vgh, histno)
                self.root.after(0, lambda: self.show_vipfsn_complete(histno, data))
            except Exception as e:
                # 這裡改成先把錯誤字串存起來，再用 lambda 的預設參數傳進去
                err_msg = str(e)
                self.root.after(
                    0,
                    lambda msg=err_msg: self.show_vipfsn_failed(msg)
                )

        threading.Thread(target=worker, daemon=True).start()


    
    def search_complete(self, count):
        """搜尋完成處理"""
        self.search_btn.config(state='normal', text="開始搜尋")
        self.progress_bar.stop()
        self.progress_var.set(f"搜尋完成，找到 {count} 筆病人資料")
        
        if count > 0:
            self.result_text.insert(tk.END, f"✅ 搜尋完成！找到 {count} 筆病人資料\n")
            self.result_text.insert(tk.END, "病人列表：\n")
            
            for i, pat in enumerate(self.pat_data[:10], 1):  # 只顯示前10筆
                patient_info = " | ".join(pat[:3]) if len(pat) >= 3 else " | ".join(pat)
                self.result_text.insert(tk.END, f"{i}. {patient_info}\n")
            
            if len(self.pat_data) > 10:
                self.result_text.insert(tk.END, f"... 還有 {len(self.pat_data) - 10} 筆資料\n")
            
            self.result_text.insert(tk.END, "\n可以點擊「產生Word報告」來生成完整報告。\n")
            self.generate_btn.config(state='normal')
        else:
            self.result_text.insert(tk.END, "❌ 未找到符合條件的病人資料\n")
        
        self.result_text.see(tk.END)   
    
    def search_failed(self, error_msg):
        """搜尋失敗處理"""
        self.search_btn.config(state='normal', text="開始搜尋")
        self.progress_bar.stop()
        self.progress_var.set("搜尋失敗")
        messagebox.showerror("搜尋失敗", f"搜尋過程發生錯誤：{error_msg}")
    
    def generate_report(self):
        """產生Word報告"""
        if not self.pat_data:
            messagebox.showerror("錯誤", "沒有病人資料可產生報告")
            return
        
        # 禁用按鈕
        self.generate_btn.config(state='disabled', text="產生中...")
        self.progress_bar.start()
        self.progress_var.set("正在產生Word報告...")
        
        def generate_thread():
            try:
                filename = self.create_word_document()
                self.root.after(0, lambda: self.generate_complete(filename))
            except Exception as e:
                self.root.after(0, lambda: self.generate_failed(str(e)))
        
        threading.Thread(target=generate_thread, daemon=True).start()
    
    def create_word_document(self):
        """創建Word文件（簡化版的原始函式）"""
        doc = Document()
        
        section = doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # 設定邊界
        section.top_margin = Pt(30)
        section.bottom_margin = Pt(30)
        section.left_margin = Pt(30)
        section.right_margin = Pt(30)
        
        style = doc.styles['Normal']
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run("日期:" + datetime.now().strftime('%Y-%m-%d') + " 醫師: " + self.docID)
        run.font.size = Pt(6)

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '病人資料'
        hdr_cells[1].text = 'Assessment'
        hdr_cells[2].text = 'Lab Data+drug'
        for cell in hdr_cells:
            set_font_size(cell, 6)
        
        for idx, pat in enumerate(self.pat_data):
            # 更新進度
            progress = f"正在處理第 {idx + 1}/{len(self.pat_data)} 位病人..."
            self.root.after(0, lambda p=progress: self.progress_var.set(p))
            
            row_cells = table.add_row().cells
            if len(pat) < 3:
                continue
            
            if self.search_type == "ward":
                ID = pat[2]
            else:
                ID = pat[1]
            
            # 這裡調用原始的 generate_table_report 函式
            try:
                self.generate_table_report_gui(doc, ID, row_cells, pat)
                for cell in row_cells:
                    set_font_size(cell, 6)
            except Exception as e:
                print(f"處理病人 {ID} 時發生錯誤: {e}")
                continue
            
            time.sleep(random.randint(3, 8))
            if idx % 10 == 0 and idx != 0:
                time.sleep(60)

        # 保存Word文件
        set_font_size(doc, 6)
        set_paragraph_spacing(doc, spacing=0)
        filename = datetime.now().strftime('%Y%m%d') + "_" + self.docID + "_" + "patient_list" + '.docx'
        doc.save(filename)
        return filename
    
    def generate_table_report_gui(self, doc, ID, row_cells, pat):
        """簡化版的表格報告生成（為了GUI穩定性）"""
        # 基本病人資訊
        info_cell = row_cells[0]
        paragraph = info_cell.paragraphs[0]
        paragraph.add_run("\n".join(pat))
        
        TPR = get_TPR(self.vgh, ID)
        
        if not TPR.empty and all(col in TPR.columns for col in ["體溫","心跳","呼吸","收縮壓","舒張壓"]):
            run = paragraph.add_run("\n")
            paragraph.add_run("\\".join(list(TPR[["體溫","心跳","呼吸","收縮壓","舒張壓"]].iloc[0])))
        
    
        TPR_img = get_TPR_img(self.vgh, ID)
        run = paragraph.add_run()
        image_path = 'downloaded_image.jpg'
        run.add_picture(image_path, width=Inches(1.5))
        os.remove(image_path)
        

        BW_BL = get_BW_BL(self.vgh, ID, adminID="all")
        if not BW_BL.empty and all(col in BW_BL.columns for col in ["身高","體重"]):
            BW_BL = BW_BL[["身高","體重"]]
            add_table(info_cell, BW_BL.head(2))
        
        # Assessment
        try:
            assessment_cell = row_cells[1]
            paragraph = assessment_cell.paragraphs[0]
            progress_note = get_progress_note(self.vgh, ID, num=5)
            if progress_note:
                for i in range(len(progress_note)):
                    assessment = progress_note[i]["Assessment"]
                    if "Ditto" not in assessment and len(assessment) >= 5:
                        assessment = assessment.replace('\r', '')
                        paragraph.add_run(assessment)
                        break
        except:
            pass
        
        # Lab Data
        Lab_cells = row_cells[2]
        try:
            patIO = get_drainage(self.vgh, ID)
            if not patIO.empty and all(col in patIO.columns for col in ["項目","總量"]):
                add_table(Lab_cells, patIO[["項目","總量"]])
                p = Lab_cells.paragraphs[-1]._element
                p.getparent().remove(p)
        except:
            pass

        
        p = Lab_cells.paragraphs[0]._element
        p.getparent().remove(p)
        report_num = 3
        report_name, recent_report = get_recent_report(self.vgh, ID, report_num=report_num)
        for i in range(len(report_name)):
            Lab_cells.add_paragraph(report_name[i])


        SMAC = get_res_report(self.vgh, ID, resdtype="SMAC")
        if not SMAC.empty and "日期" in SMAC.columns:
            SMAC["日期"] = SMAC["日期"].apply(convert_date)
            required_cols = ["日期","NA","K","BUN","CREA","ALT","BILIT","CRP"]
            if all(col in SMAC.columns for col in required_cols):
                SMAC = SMAC[required_cols]
                SMAC = SMAC.loc[~(SMAC[required_cols] == '-').all(axis=1)]
                add_table(Lab_cells, SMAC.tail(3))
                p = Lab_cells.paragraphs[-1]._element
                p.getparent().remove(p)
        

        CBC = get_res_report(self.vgh, ID, resdtype="CBC")
        if not CBC.empty and "日期" in CBC.columns:
            CBC["日期"] = CBC["日期"].apply(convert_date)
            required_cols = ["日期","WBC","HGB","PLT",'SEG', 'PT', 'APTT']
            if all(col in CBC.columns for col in required_cols):
                CBC = CBC[required_cols]
                CBC = CBC.loc[~(CBC[required_cols] == '-').all(axis=1)]
                add_table(Lab_cells, CBC.tail(3))
                p = Lab_cells.paragraphs[-1]._element
                p.getparent().remove(p)

        drug = get_drug(self.vgh, ID)
        if not drug.empty and all(col in drug.columns for col in ["學名","開始日","狀態"]):
            drug["學名"] = drug["學名"].apply(convert_drug)
            drug["開始日"] = drug["開始日"].apply(convert_drug_date)
            required_cols = ["學名","劑量","途徑","頻次","開始日"]
            if all(col in drug.columns for col in required_cols):
                add_table(Lab_cells, drug[drug["狀態"]=="使用中"][required_cols])
    
    def generate_complete(self, filename):
        """報告產生完成處理"""
        self.generate_btn.config(state='normal', text="產生Word報告")
        self.progress_bar.stop()
        self.progress_var.set("Word報告產生完成")
        
        self.result_text.insert(tk.END, f"\n✅ Word報告產生完成！\n")
        self.result_text.insert(tk.END, f"檔案名稱: {filename}\n")
        self.result_text.see(tk.END)
        
        messagebox.showinfo("完成", f"Word報告已成功產生！\n檔案名稱: {filename}")
    
    def generate_failed(self, error_msg):
        """報告產生失敗處理"""
        self.generate_btn.config(state='normal', text="產生Word報告")
        self.progress_bar.stop()
        self.progress_var.set("報告產生失敗")
        
        self.result_text.insert(tk.END, f"\n❌ Word報告產生失敗: {error_msg}\n")
        self.result_text.see(tk.END)
        
        messagebox.showerror("錯誤", f"Word報告產生失敗：{error_msg}")
    
    def clear_results(self):
        """清除結果"""
        self.result_text.delete(1.0, tk.END)
        self.pat_data = []
        self.generate_btn.config(state='disabled')
        self.progress_var.set("等待開始...")


def main():
    """主程式入口"""
    root = tk.Tk()
    app = VGHPatientGUI(root)
    
    # 設定窗口圖標（如果有的話）
    try:
        root.iconbitmap('icon.ico')  # 如果有圖標文件
    except:
        pass
    
    # 設定窗口關閉事件
    def on_closing():
        if messagebox.askokcancel("退出", "確定要退出程式嗎？"):
            root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    # 啟動GUI
    root.mainloop()


if __name__ == "__main__":
    main()
