#!/usr/bin/env python
# coding: utf-8

import requests
from bs4 import BeautifulSoup
import urllib.parse
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from PIL import Image
from docx.oxml.ns import qn
import os
from VGH_function import *
from datetime import datetime, timedelta
import pwinput
import platform
import time
import random


class VGHLogin:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        self.csrf_token = None
        self.base_url = "https://eip.vghtpe.gov.tw/login.php"
    
    def get_login_page(self):
        """取得登入頁面並解析CSRF token"""
        try:
            response = self.session.get(self.base_url)
            response.raise_for_status()
            
            # 解析HTML取得CSRF token
            soup = BeautifulSoup(response.text, 'html.parser')
            csrf_meta = soup.find('meta', {'name': 'csrf-token'})
            if csrf_meta:
                self.csrf_token = csrf_meta.get('content')
                # print(f"取得CSRF Token: {self.csrf_token}")
            
            return True
        except requests.RequestException as e:
            print(f"取得登入頁面失敗: {e}")
            return False
    
    def login(self, username, password):
        """執行登入"""
        if not self.get_login_page():
            return False
        
        # 準備登入資料
        login_data = {
            'login_name': username,
            'password': password,
            'loginCheck': '1',
            'fromAjax': '1'
        }
        
        # 設定headers
        headers = {
            'X-CSRF-TOKEN': self.csrf_token,
            'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
            'X-Requested-With': 'XMLHttpRequest',
            'Referer': self.base_url
        }
        
        try:
            # 發送登入請求
            login_url = urllib.parse.urljoin(self.base_url, '/login_action.php')
            response = self.session.post(
                login_url,
                data=login_data,
                headers=headers
            )
            response.raise_for_status()
            
            # 解析回應
            result = response.json()

            if 'error' in result:
                error_code = int(result['error'])
                if error_code == 0:
                    print("登入成功!")
                    if 'url' in result:
                        print(f"重定向至: {result['url']}")
                        # 可以選擇跟隨重定向
                        dashboard_response = self.session.get("https://eip.vghtpe.gov.tw/"+result['url'])
                        login_url="https://eip.vghtpe.gov.tw/"+dashboard_response.text.split("/")[1][:-2]
                        dashboard_response = self.session.get(login_url)
                        return True
                else:
                    print(f"登入失敗: {result.get('err_msg', '未知錯誤')}")
                    return False
            else:
                print("未預期的回應格式")
                return False
                
        except requests.RequestException as e:
            print(f"登入請求失敗: {e}")
            return False
        except ValueError as e:
            print(f"JSON解析失敗: {e}")
            return False
    
    def get_page_after_login(self, url):
        """登入後取得其他頁面"""
        try:
            response = self.session.get(url)
            response.raise_for_status()
            return response.text
        except requests.RequestException as e:
            print(f"取得頁面失敗: {e}")
            return None

    def get_img_after_login(self, url):
        """登入後取得其他頁面"""
        try:
            response = self.session.get(url)
            response.raise_for_status()
            return response
        except requests.RequestException as e:
            print(f"取得頁面失敗: {e}")
            return None


def clear_terminal():
    if platform.system() == "Windows":
        os.system('cls')
    else:
        os.system('clear')


def login_to_system():
    """處理登入邏輯"""
    clear_terminal()
    print("""
    此程式可以自動查詢病人資料，製作一份WORD的查房摘要
    請先打入入口網帳號密碼，之後輸入燈號或式直接按ENTER之後輸入病房
    ***注意 若病人太多可能會被資訊室鎖住該台電腦一陣子，可以重新開機後稍等一下
    作者的燈號為: 8375K，如果有任何問題或建議，歡迎聯絡!!!
    """)
    
    while True:
        # 要求使用者輸入帳號與密碼
        username = input("請輸入帳號：")
        password = pwinput.pwinput(prompt='密碼: ', mask='*')
        
        try:
            # 使用 VGHLogin 進行登入
            vgh = VGHLogin()
            if vgh.login(username, password):
                print("✅ 登入成功！")
                clear_terminal()
                return vgh
        except Exception as e:
            clear_terminal()
            print("""
    此程式可以自動查詢病人資料，製作一份WORD的查房摘要
    請先打入入口網帳號密碼，之後輸入燈號或式直接按ENTER之後輸入病房
    ***注意 若病人太多可能會被資訊室鎖住該台電腦一陣子，可以重新開機後稍等一下
    作者的燈號為: 8375K，如果有任何問題或建議，歡迎聯絡!!!
    """)
            print(f"⚠️ 登入失敗: {e}，請重新輸入帳號與密碼。\n")


def get_search_type(vgh):
    """取得搜尋類型"""
    Search_type = input("選擇要如何搜尋病人，依燈號請輸入doc, 病房請出入ward, 病歷號請輸入 pat:")
    
    while not Search_type=="doc" and not Search_type=="ward" and not Search_type=="pat":
        clear_terminal()
        print("輸入錯誤，請重新輸入")
        Search_type=input("選擇要如何搜尋病人，依燈號請輸入doc, 病房請出入ward, 病歷號請輸入 pat:")
    docID=""
    if Search_type=="doc":
        ward="0"
        docID=input("請輸入燈號(四碼):")
        pat_data=get_searched_patient(vgh,ward=ward,patID="",docID=docID)
        return(pat_data,docID, Search_type)
    elif Search_type=="ward":
        docID=""
        ward=input("請輸入病房(Ex A101):")
        pat_data=get_searched_patient(vgh,ward=ward,patID="",docID="")
        return(pat_data,docID, Search_type)
    else:
        ward="0"
        docID=""
        pat_data=[]
        patID=input("請輸入病歷號(若不須再輸入請直接按enter):")
        while not patID=="":
            pat_data.append(get_searched_patient(vgh,ward=ward,patID=patID,docID="")[0])
            patID=input("請輸入病歷號(若不須再輸入請直接按enter):")
        return(pat_data,docID, Search_type)


def set_paragraph_spacing(doc, spacing=0):
    """Set paragraph spacing for all paragraphs in the document."""
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = Pt(spacing)
        paragraph.paragraph_format.space_before = Pt(spacing)
        paragraph.paragraph_format.space_after = Pt(spacing)


def set_font_size(doc, size):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)


def add_table(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # 設置表頭的字體大小
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = str(column_name)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(6)
            paragraph.paragraph_format.line_spacing = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    last_paragraph = doc.paragraphs[-1]
    last_paragraph.paragraph_format.space_after = Pt(0)
    last_paragraph.paragraph_format.space_before = Pt(0)
    last_paragraph.paragraph_format.line_spacing = Pt(0)
    
    # 添加數據行
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6)
                paragraph.paragraph_format.line_spacing = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                
    for col in table.columns:
        max_length = max(len(cell.text) for cell in col.cells)
        col_width = Inches(max_length)
        for cell in col.cells:
            cell.width = col_width


def convert_date(date_str):
    date_str = date_str[3:8]
    return date_str


def convert_drug(data_drug):
    data_drug = data_drug.split(" ")[:2]
    data_drug = " ".join(data_drug)
    return data_drug


def convert_drug_date(data_drug_date):
    data_drug_date = data_drug_date[5:10]
    return data_drug_date


def generate_table_report(vgh, doc, ID, row_cells, pat):
    """生成表格報告（使用 VGHLogin 而非 driver）"""
    print(ID)
    
    info_cell = row_cells[0]
    paragraph = info_cell.paragraphs[0]
    paragraph.add_run("\n".join(pat))
    try:
        TPR = get_TPR(vgh, ID)
        if not TPR.empty and all(col in TPR.columns for col in ["體溫","心跳","呼吸","收縮壓","舒張壓"]):
            run = paragraph.add_run("\n")
            paragraph.add_run("\\".join(list(TPR[["體溫","心跳","呼吸","收縮壓","舒張壓"]].iloc[0])))
    except:
        pass
    
    try:
        TPR_img = get_TPR_img(vgh, ID)
        run = paragraph.add_run()
        image_path = 'downloaded_image.jpg'
        run.add_picture(image_path, width=Inches(1.5))
        os.remove(image_path)
    except:
        pass

    try:
        BW_BL = get_BW_BL(vgh, ID, adminID="all")
        if not BW_BL.empty and all(col in BW_BL.columns for col in ["身高","體重"]):
            BW_BL = BW_BL[["身高","體重"]]
            add_table(info_cell, BW_BL.head(2))
    except:
        pass
    try:
        assessment_cell = row_cells[1]
        paragraph = assessment_cell.paragraphs[0]
        progress_note = get_progress_note(vgh, ID, num=5)
        for i in range(len(progress_note)):
            assessment = progress_note[i]["Assessment"]
            if "Ditto" in assessment or len(assessment) < 5:
                pass
            else:
                break
        assessment = assessment.replace('\r', '')
        paragraph.add_run(assessment)
    except:
        pass

    Lab_cells = row_cells[2]
    
    try:
        patIO = get_drainage(vgh, ID)
        if not patIO.empty and all(col in patIO.columns for col in ["項目","總量"]):
            add_table(Lab_cells, patIO[["項目","總量"]])
            p = Lab_cells.paragraphs[-1]._element
            p.getparent().remove(p)
    except:
        pass
    
    try:
        p = Lab_cells.paragraphs[0]._element
        p.getparent().remove(p)
        report_num = 3
        report_name, recent_report = get_recent_report(vgh, ID, report_num=report_num)
        for i in range(len(report_name)):
            Lab_cells.add_paragraph(report_name[i])
    except:
        pass

    try:
        SMAC = get_res_report(vgh, ID, resdtype="SMAC")
        if not SMAC.empty and "日期" in SMAC.columns:
            SMAC["日期"] = SMAC["日期"].apply(convert_date)
            required_cols = ["日期","NA","K","BUN","CREA","ALT","BILIT","CRP"]
            if all(col in SMAC.columns for col in required_cols):
                SMAC = SMAC[required_cols]
                SMAC = SMAC.loc[~(SMAC[required_cols] == '-').all(axis=1)]
                add_table(Lab_cells, SMAC.tail(3))
                p = Lab_cells.paragraphs[-1]._element
                p.getparent().remove(p)
    except:
        pass

    try:
        CBC = get_res_report(vgh, ID, resdtype="CBC")
        if not CBC.empty and "日期" in CBC.columns:
            CBC["日期"] = CBC["日期"].apply(convert_date)
            required_cols = ["日期","WBC","HGB","PLT",'SEG', 'PT', 'APTT']
            if all(col in CBC.columns for col in required_cols):
                CBC = CBC[required_cols]
                CBC = CBC.loc[~(CBC[required_cols] == '-').all(axis=1)]
                add_table(Lab_cells, CBC.tail(3))
                p = Lab_cells.paragraphs[-1]._element
                p.getparent().remove(p)
    except:
        pass

    try:
        drug = get_drug(vgh, ID)
        if not drug.empty and all(col in drug.columns for col in ["學名","開始日","狀態"]):
            drug["學名"] = drug["學名"].apply(convert_drug)
            drug["開始日"] = drug["開始日"].apply(convert_drug_date)
            required_cols = ["學名","劑量","途徑","頻次","開始日"]
            if all(col in drug.columns for col in required_cols):
                add_table(Lab_cells, drug[drug["狀態"]=="使用中"][required_cols])
    except:
        pass


def create_word_document(pat_data, docID, search_type, vgh):
    """創建 Word 文件"""
    doc = Document()
    
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # 設定邊界
    section.top_margin = Pt(30)
    section.bottom_margin = Pt(30)
    section.left_margin = Pt(30)
    section.right_margin = Pt(30)
    
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run("日期:" + datetime.now().strftime('%Y-%m-%d') + " 醫師: " + docID)
    run.font.size = Pt(6)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '病人資料'
    hdr_cells[1].text = 'Assessment'
    hdr_cells[2].text = 'Lab Data+drug'
    for cell in hdr_cells:
        set_font_size(cell, 6)

    for idx, pat in enumerate(pat_data):
        row_cells = table.add_row().cells
        if len(pat) < 3:
            continue
        if search_type == "ward":
            ID = pat[2]
        else:
            ID = pat[1]
        generate_table_report(vgh=vgh, doc=doc, ID=ID, row_cells=row_cells, pat=pat)
        for cell in row_cells:
            set_font_size(cell, 6)
        time.sleep(random.randint(3, 8))
        if idx % 10 == 0 and idx != 0:
            print("wait a while")
            time.sleep(60)

    for idx, col in enumerate(table.columns):
        max_length = max(len(cell.text) for cell in col.cells)
        col_width = Inches(max_length)
        if idx == 2:
            col_width = Inches(max_length)
        for cell in col.cells:
            cell.width = col_width

    # 設置所有文本字體為 6 號
    set_font_size(doc, 6)
    set_paragraph_spacing(doc, spacing=0)

    # 保存 Word 文件
    filename = datetime.now().strftime('%Y%m%d') + "_" + docID + "_" + "patient_list" + '.docx'
    doc.save(filename)
    print("儲存為" + filename)


def main():
    """主程式"""
    # 登入系統
    vgh = login_to_system()
    
    time.sleep(0.5)
    
    # 取得搜尋類型
    
    pat_data, docID, search_type = get_search_type(vgh)
    
    # 創建 Word 文件
    create_word_document(pat_data, docID,search_type, vgh)
    
    print("程式執行完成！")


if __name__ == "__main__":
    main()